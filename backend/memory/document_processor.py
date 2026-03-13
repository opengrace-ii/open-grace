"""
Document Processor - Extract text from various document formats.

Supports:
- PDF (.pdf)
- Word Documents (.docx, .doc)
- Excel Spreadsheets (.xls, .xlsx, .csv)
- Text files (.txt, .md, .rst)
- Code files (.py, .js, .ts, etc.)
"""

import os
import csv
import io
from typing import Optional, Dict, Any, List
from pathlib import Path
from dataclasses import dataclass

from backend.observability.logger import get_logger


@dataclass
class ExtractedContent:
    """Result of document extraction."""
    text: str
    metadata: Dict[str, Any]
    pages: Optional[int] = None
    sheets: Optional[List[str]] = None


class DocumentProcessor:
    """
    Process various document formats and extract text content.
    
    Usage:
        processor = DocumentProcessor()
        
        # Extract from PDF
        content = processor.extract("document.pdf")
        print(content.text)
        
        # Extract from Excel
        content = processor.extract("data.xlsx")
        print(content.text)  # All sheets combined
        print(content.sheets)  # List of sheet names
    """
    
    def __init__(self):
        self.logger = get_logger()
        self._parsers = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
            '.csv': self._extract_csv,
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.rst': self._extract_text,
            '.json': self._extract_text,
            '.xml': self._extract_text,
            '.html': self._extract_html,
            '.htm': self._extract_html,
        }
    
    def extract(self, file_path: str) -> ExtractedContent:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            ExtractedContent with text and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        # Get the appropriate parser
        parser = self._parsers.get(extension, self._extract_text)
        
        try:
            content = parser(file_path)
            
            # Add common metadata
            content.metadata.update({
                'filename': path.name,
                'file_type': extension,
                'file_size': path.stat().st_size,
                'source': str(path.absolute())
            })
            
            return content
            
        except Exception as e:
            self.logger.error(f"Failed to extract {file_path}: {e}")
            # Fallback to text extraction
            return self._extract_text(file_path)
    
    def _extract_pdf(self, file_path: str) -> ExtractedContent:
        """Extract text from PDF."""
        try:
            import PyPDF2
            
            text_parts = []
            metadata = {}
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                
                metadata['pages'] = num_pages
                metadata['pdf_title'] = reader.metadata.get('/Title', '') if reader.metadata else ''
                metadata['pdf_author'] = reader.metadata.get('/Author', '') if reader.metadata else ''
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"\n--- Page {i + 1} ---\n{page_text}")
                    except Exception as e:
                        self.logger.warning(f"Could not extract page {i + 1}: {e}")
                
                full_text = '\n'.join(text_parts)
                
                return ExtractedContent(
                    text=full_text,
                    metadata=metadata,
                    pages=num_pages
                )
                
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
            return self._extract_text(file_path)
    
    def _extract_docx(self, file_path: str) -> ExtractedContent:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- Table {i + 1} ---")
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            # Get metadata
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
            }
            
            if doc.core_properties:
                metadata['title'] = doc.core_properties.title or ''
                metadata['author'] = doc.core_properties.author or ''
            
            return ExtractedContent(
                text='\n'.join(text_parts),
                metadata=metadata
            )
            
        except ImportError:
            self.logger.warning("python-docx not installed. Install with: pip install python-docx")
            return self._extract_text(file_path)
    
    def _extract_doc(self, file_path: str) -> ExtractedContent:
        """Extract text from older DOC format."""
        # Try using LibreOffice conversion first
        try:
            return self._convert_with_libreoffice(file_path)
        except Exception as e:
            self.logger.warning(f"LibreOffice conversion failed: {e}")
            
        # Fallback: try antiword if available
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return ExtractedContent(
                    text=result.stdout,
                    metadata={'extractor': 'antiword'}
                )
        except (subprocess.SubprocessError, FileNotFoundError):
            pass
        
        # Final fallback
        return self._extract_text(file_path)
    
    def _extract_excel(self, file_path: str) -> ExtractedContent:
        """Extract text from Excel files."""
        try:
            import pandas as pd
            
            # Read all sheets
            xls = pd.ExcelFile(file_path)
            sheet_names = xls.sheet_names
            
            text_parts = []
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                text_parts.append(f"\n=== Sheet: {sheet_name} ===")
                text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                text_parts.append("")
                
                # Convert DataFrame to text
                # Option 1: Include headers
                text_parts.append(" | ".join(str(col) for col in df.columns))
                text_parts.append("-" * 50)
                
                # Add rows (limit to avoid huge outputs)
                max_rows = 1000
                for idx, row in df.head(max_rows).iterrows():
                    row_text = " | ".join(str(val) for val in row.values)
                    text_parts.append(row_text)
                
                if len(df) > max_rows:
                    text_parts.append(f"\n... ({len(df) - max_rows} more rows) ...")
                
                text_parts.append("")
            
            metadata = {
                'sheets': sheet_names,
                'total_sheets': len(sheet_names)
            }
            
            return ExtractedContent(
                text='\n'.join(text_parts),
                metadata=metadata,
                sheets=sheet_names
            )
            
        except ImportError:
            self.logger.warning("pandas not installed. Install with: pip install pandas openpyxl")
            return self._extract_text(file_path)
    
    def _extract_csv(self, file_path: str) -> ExtractedContent:
        """Extract text from CSV."""
        try:
            text_parts = []
            row_count = 0
            
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                
                # Get headers
                try:
                    headers = next(reader)
                    text_parts.append(" | ".join(headers))
                    text_parts.append("-" * 50)
                    row_count += 1
                except StopIteration:
                    pass
                
                # Get rows (limit to avoid huge outputs)
                max_rows = 1000
                for i, row in enumerate(reader):
                    if i >= max_rows:
                        break
                    text_parts.append(" | ".join(row))
                    row_count += 1
                
                # Count remaining rows
                remaining = sum(1 for _ in reader)
                if remaining > 0:
                    text_parts.append(f"\n... ({remaining} more rows) ...")
            
            metadata = {
                'rows': row_count + remaining if 'remaining' in dir() else row_count,
                'columns': len(headers) if 'headers' in dir() else 0
            }
            
            return ExtractedContent(
                text='\n'.join(text_parts),
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"CSV extraction failed: {e}")
            return self._extract_text(file_path)
    
    def _extract_text(self, file_path: str) -> ExtractedContent:
        """Extract text from plain text files."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    
                return ExtractedContent(
                    text=text,
                    metadata={'encoding': encoding}
                )
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with replacement
        with open(file_path, 'rb') as f:
            binary = f.read()
            text = binary.decode('utf-8', errors='replace')
            
        return ExtractedContent(
            text=text,
            metadata={'encoding': 'utf-8-with-replacement'}
        )
    
    def _extract_html(self, file_path: str) -> ExtractedContent:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text(separator='\n', strip=True)
            
            # Clean up whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)
            
            # Get title
            title = soup.title.string if soup.title else ''
            
            return ExtractedContent(
                text=text,
                metadata={
                    'title': title,
                    'extractor': 'beautifulsoup'
                }
            )
            
        except ImportError:
            return self._extract_text(file_path)
    
    def _convert_with_libreoffice(self, file_path: str) -> ExtractedContent:
        """Convert document using LibreOffice and extract text."""
        import subprocess
        import tempfile
        import shutil
        
        with tempfile.TemporaryDirectory() as tmpdir:
            # Convert to text using LibreOffice
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'txt:Text',
                    '--outdir', tmpdir,
                    file_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            # Find the converted file
            path = Path(file_path)
            txt_file = Path(tmpdir) / f"{path.stem}.txt"
            
            if not txt_file.exists():
                raise FileNotFoundError("Converted file not found")
            
            # Read the text
            with open(txt_file, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return ExtractedContent(
                text=text,
                metadata={'extractor': 'libreoffice'}
            )
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self._parsers.keys())
    
    def can_process(self, file_path: str) -> bool:
        """Check if file can be processed."""
        path = Path(file_path)
        return path.suffix.lower() in self._parsers or self._is_text_file(file_path)
    
    def _is_text_file(self, file_path: str) -> bool:
        """Check if file is a text file by examining content."""
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                # Check for null bytes (binary file indicator)
                return b'\x00' not in chunk
        except:
            return False


# Global processor instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the global document processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor